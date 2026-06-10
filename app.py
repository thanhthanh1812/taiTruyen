import os
import re
import time
import queue
import uuid
import threading
from urllib.parse import urlparse, parse_qs
import requests
from bs4 import BeautifulSoup
from docx import Document
from flask import Flask, render_template, request, jsonify, Response, send_from_directory

app = Flask(__name__)

# Thư mục lưu file Word tạm thời
DOWNLOAD_FOLDER = os.path.join(app.root_path, 'static', 'downloads')
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# Lưu trữ các hàng đợi log và tên file tương ứng với mỗi tác vụ
# Key: task_id (str), Value: Queue (log messages)
task_queues = {}
# Key: task_id (str), Value: filename (str)
task_files = {}

class TruyenHdaCrawler:
    def __init__(self, cookies_str, log_func):
        self.base_url = "https://truyenhdc.com"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Cookie': cookies_str.strip()
        }
        self.doc = Document()
        self.log = log_func

    def get_soup(self, url):
        try:
            response = requests.get(url, headers=self.headers, timeout=15)
            if response.status_code == 200:
                return BeautifulSoup(response.text, 'html.parser')
            else:
                self.log(f"[!] Lỗi HTTP {response.status_code} khi truy cập: {url}")
                return None
        except Exception as e:
            self.log(f"[!] Lỗi kết nối: {str(e)}")
            return None

    def get_chapter_ids(self, story_id, n_pages):
        chapter_ids = []
        self.log(f"[*] Đang quét danh sách chương (ID Truyện: {story_id})...")

        for page in range(1, n_pages + 1):
            url = f"{self.base_url}/user/quan-ly-truyen/dsc/?id={story_id}&n={page}"
            self.log(f"    -> Đang quét trang danh sách chương {page}/{n_pages}...")

            soup = self.get_soup(url)
            if not soup:
                continue

            list_group = soup.find('div', class_='list-group', id='dsc')
            if not list_group:
                self.log("    [!] Không tìm thấy danh sách chương. Hãy kiểm tra lại Cookie hoặc ID truyện.")
                break

            items = list_group.find_all('div', class_='list-group-item')

            for item in items:
                # Tìm thẻ a chỉnh sửa chương
                edit_btn = item.find('a', href=lambda x: x and 'edit-chuong' in x)
                if edit_btn:
                    href = edit_btn.get('href')
                    parsed_url = urlparse(href)
                    query_params = parse_qs(parsed_url.query)

                    if 'id' in query_params:
                        c_id = query_params['id'][0]
                        chapter_ids.append(c_id)

        # Đảo ngược thứ tự để lấy chương cũ trước, chương mới sau (đọc xuôi)
        chapter_ids.reverse()
        self.log(f"[*] Tìm thấy tổng cộng {len(chapter_ids)} chương.")
        return chapter_ids

    def get_chapter_content(self, chapter_id):
        url = f"{self.base_url}/user/quan-ly-truyen/edit-chuong/?id={chapter_id}"
        soup = self.get_soup(url)

        if not soup:
            return None, None

        # Lấy tên chương
        title_input = soup.find('input', id='ten_chuong')
        title = title_input.get('value') if title_input else f"Chương {chapter_id}"

        # Lấy nội dung
        content_div = soup.find('div', id='richeditor')
        if content_div:
            # Xử lý xuống dòng phù hợp
            for br in content_div.find_all("br"):
                br.replace_with("\n")
            for p in content_div.find_all("p"):
                p.append("\n")

            content = content_div.get_text()
            content = re.sub(r'\n\s*\n', '\n\n', content).strip()
        else:
            content = "[Không lấy được nội dung chương này]"

        return title, content

    def save_to_file(self, chapter_ids, output_path):
        self.log("[*] Bắt đầu tải nội dung từng chương...")
        total = len(chapter_ids)
        for idx, c_id in enumerate(chapter_ids):
            title, content = self.get_chapter_content(c_id)

            if title and content:
                self.log(f"    [{idx + 1}/{total}] Đã tải thành công: {title}")
                
                # Định dạng tên chương viết hoa và canh giữa
                heading = self.doc.add_heading(title.upper(), level=1)
                heading.alignment = 1
                
                # Thêm paragraph nội dung
                self.doc.add_paragraph(content)
                self.doc.add_page_break()
            else:
                self.log(f"    [!] Gặp lỗi khi tải chương ID: {c_id}")

            # Tránh spam request
            time.sleep(0.5)

        self.doc.save(output_path)
        self.log("[*] Hoàn tất việc ghi file Word!")


def run_crawl_thread(task_id, cookies_str, story_id, total_pages, output_filename):
    q = task_queues[task_id]
    
    def log_func(msg):
        q.put(msg)
        
    try:
        # Chuẩn hóa tên file đầu ra (.docx)
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'
        
        stored_filename = f"{task_id}_{output_filename}"
        file_path = os.path.join(DOWNLOAD_FOLDER, stored_filename)
        task_files[task_id] = stored_filename
        
        crawler = TruyenHdaCrawler(cookies_str, log_func)
        ids = crawler.get_chapter_ids(story_id, total_pages)
        if ids:
            crawler.save_to_file(ids, file_path)
            q.put(f"[DONE] {output_filename}")
        else:
            q.put("[ERROR] Không tìm thấy chương nào để tải. Vui lòng kiểm tra lại Cookie và ID truyện.")
    except Exception as e:
        q.put(f"[ERROR] Đã xảy ra lỗi trong quá trình tải: {str(e)}")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/start-download', methods=['POST'])
def start_download():
    data = request.json or {}
    cookie_name = data.get('cookie_name', '').strip()
    cookie_value = data.get('cookie_value', '').strip()
    story_id = data.get('story_id', '').strip()
    
    try:
        total_pages = int(data.get('total_pages', 1))
    except ValueError:
        total_pages = 1
        
    output_filename = data.get('output_filename', '').strip()
    if not output_filename:
        output_filename = f"truyen_{story_id}.docx"
        
    if not cookie_name or not cookie_value or not story_id:
        return jsonify({
            "status": "error",
            "message": "Vui lòng nhập đầy đủ các trường bắt buộc (Cookie, ID Truyện)."
        }), 400
        
    # Tạo chuỗi cookie hoàn chỉnh
    cookies_str = f"{cookie_name}={cookie_value}"
    
    # Tạo task_id duy nhất
    task_id = str(uuid.uuid4())
    task_queues[task_id] = queue.Queue()
    
    # Chạy thread crawler không đồng bộ
    t = threading.Thread(
        target=run_crawl_thread,
        args=(task_id, cookies_str, story_id, total_pages, output_filename),
        daemon=True
    )
    t.start()
    
    return jsonify({
        "status": "started",
        "task_id": task_id
    })


@app.route('/api/stream-logs/<task_id>')
def stream_logs(task_id):
    def event_stream():
        q = task_queues.get(task_id)
        if not q:
            yield "data: [ERROR] Không tìm thấy tiến trình tải tương ứng.\n\n"
            return
        
        while True:
            try:
                # Đợi tối đa 30s để tránh ngắt kết nối
                msg = q.get(timeout=30)
                if msg.startswith("[DONE]"):
                    filename = msg[7:]
                    yield f"event: complete\ndata: {filename}\n\n"
                    break
                elif msg.startswith("[ERROR]"):
                    error_msg = msg[8:]
                    yield f"event: error\ndata: {error_msg}\n\n"
                    break
                else:
                    yield f"data: {msg}\n\n"
            except queue.Empty:
                # Gửi heartbeat định kỳ để giữ kết nối SSE không bị đóng
                yield "data: [SYSTEM] Đang chờ phản hồi từ website...\n\n"
            except Exception as e:
                yield f"event: error\ndata: {str(e)}\n\n"
                break
                
    return Response(event_stream(), mimetype="text/event-stream")


@app.route('/api/download/<task_id>')
def download_file(task_id):
    filename = task_files.get(task_id)
    if not filename or not os.path.exists(os.path.join(DOWNLOAD_FOLDER, filename)):
        return "Lỗi: File tải xuống không tồn tại hoặc đã bị xóa do hết hạn.", 404
        
    # Tên file sạch hiển thị khi tải xuống (cắt bỏ phần uuid và dấu gạch dưới)
    clean_name = filename[37:]
    
    return send_from_directory(
        DOWNLOAD_FOLDER,
        filename,
        as_attachment=True,
        download_name=clean_name
    )


if __name__ == '__main__':
    # Chạy cục bộ
    print("[*] Server đang khởi chạy tại http://127.0.0.1:5000")
    app.run(host='127.0.0.1', port=5000, debug=True)

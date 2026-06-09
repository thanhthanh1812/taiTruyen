import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import time
import re
from urllib.parse import urlparse, parse_qs


class TruyenHdaCrawler:
    def __init__(self, cookies_str):
        self.base_url = "https://truyenhdc.com"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Cookie': cookies_str.strip()
        }
        self.doc = Document()

    def get_soup(self, url):
        """Hàm helper để lấy BeautifulSoup object từ URL"""
        try:
            response = requests.get(url, headers=self.headers, timeout=10)
            if response.status_code == 200:
                return BeautifulSoup(response.text, 'html.parser')
            else:
                print(f"[!] Lỗi {response.status_code} khi truy cập: {url}")
                return None
        except Exception as e:
            print(f"[!] Exception: {e}")
            return None

    def get_chapter_ids(self, story_id, n_pages):
        """Lấy danh sách ID chương từ trang quản lý danh sách chương"""
        chapter_ids = []
        print(f"[*] Đang quét danh sách chương (ID Truyện: {story_id})...")

        # Duyệt qua các trang phân trang (n)
        for page in range(1, n_pages + 1):
            url = f"{self.base_url}/user/quan-ly-truyen/dsc/?id={story_id}&n={page}"
            print(f"    -> Đang quét trang {page}/{n_pages}...")

            soup = self.get_soup(url)
            if not soup: continue

            # Tìm list group
            list_group = soup.find('div', class_='list-group', id='dsc')
            if not list_group:
                print("    [!] Không tìm thấy danh sách chương. Kiểm tra lại Cookie hoặc ID truyện.")
                break

            items = list_group.find_all('div', class_='list-group-item')

            for item in items:
                # Tìm nút Edit để lấy ID chương
                # HTML mẫu: <a ... href="/user/quan-ly-truyen/edit-chuong/?id=15282491#h2">
                edit_btn = item.find('a', href=lambda x: x and 'edit-chuong' in x)

                if edit_btn:
                    href = edit_btn.get('href')
                    parsed_url = urlparse(href)
                    query_params = parse_qs(parsed_url.query)

                    if 'id' in query_params:
                        c_id = query_params['id'][0]
                        chapter_ids.append(c_id)

        # Mặc định web hiển thị mới nhất trước, nên đảo ngược để lấy cũ nhất trước (đọc xuôi)
        # Nếu muốn giữ nguyên thứ tự hiển thị của web thì bỏ dòng này
        chapter_ids.reverse()

        print(f"[*] Tìm thấy tổng cộng {len(chapter_ids)} chương.")
        return chapter_ids

    def get_chapter_content(self, chapter_id):
        """Lấy nội dung chi tiết từ trang sửa chương"""
        url = f"{self.base_url}/user/quan-ly-truyen/edit-chuong/?id={chapter_id}"
        soup = self.get_soup(url)

        if not soup: return None, None

        # 1. Lấy tên chương từ input value
        # <input type="text" class="form-control" id="ten_chuong" ... value='Chương 1.1'>
        title_input = soup.find('input', id='ten_chuong')
        title = title_input.get('value') if title_input else f"Chapter {chapter_id}"

        # 2. Lấy nội dung từ div richeditor
        # <div id="richeditor" contenteditable="true">...</div>
        content_div = soup.find('div', id='richeditor')

        if content_div:
            # Xử lý xuống dòng: Thay thẻ <br> và <p> bằng newline để giữ format
            for br in content_div.find_all("br"):
                br.replace_with("\n")
            for p in content_div.find_all("p"):
                p.append("\n")

            content = content_div.get_text()
            # Clean up: Xóa các dòng trống thừa
            content = re.sub(r'\n\s*\n', '\n\n', content).strip()
        else:
            content = "[Không lấy được nội dung]"

        return title, content

    def save_to_file(self, chapter_ids, output_filename):
        """Duyệt qua ID, lấy nội dung và lưu vào Word"""
        print(f"[*] Bắt đầu tải nội dung và ghi vào {output_filename}...")

        for idx, c_id in enumerate(chapter_ids):
            title, content = self.get_chapter_content(c_id)

            if title and content:
                print(f"    [{idx + 1}/{len(chapter_ids)}] Đã tải: {title}")

                # Format: TÊN CHƯƠNG VIẾT HOA
                heading = self.doc.add_heading(title.upper(), level=1)
                heading.alignment = 1  # Center align (tùy chọn)

                # Format: Nội dung chương
                self.doc.add_paragraph(content)

                # Thêm ngắt trang giữa các chương
                self.doc.add_page_break()
            else:
                print(f"    [!] Lỗi tải chương ID: {c_id}")

            # Sleep nhẹ để tránh bị chặn (rate limit)
            time.sleep(0.5)

        self.doc.save(output_filename)
        print("[*] Hoàn tất!")


# ==========================================
# CẤU HÌNH VÀ CHẠY
# ==========================================

if __name__ == "__main__":
    # 1. Lấy chuỗi Cookie từ trình duyệt (F12 -> Network -> Reload trang -> Request Headers -> Cookie)
    # Copy toàn bộ chuỗi cookie và paste vào đây:
    YOUR_COOKIE = "wordpress_logged_in_fc056e1f7bf1317b0c2c528e1dc3e0ac=113380434164263925047%7C1781532530%7CB2H8fnoo9hSSLQoRuHvbzulM1T2rhncptqYB1bKCUBa%7Cc178ad3bdc4488abcdf13b56c8e6adaaca17d067a8b910243fb966d967289eca"

    # 2. ID truyện (Lấy từ URL dsc/?id=...)
    STORY_ID = "16756956"

    # 3. Số trang phân trang muốn quét (biến n)
    TOTAL_PAGES = 2

    OUTPUT_FILE = r"H\Quyến Rũ Ba Ba Của Bạn Thân.docx"

    # Chạy
    crawler = TruyenHdaCrawler(YOUR_COOKIE)

    # Bước 1: Lấy list ID
    ids = crawler.get_chapter_ids(STORY_ID, TOTAL_PAGES)

    # Bước 2: Tải nội dung và lưu
    if ids:
        crawler.save_to_file(ids, OUTPUT_FILE)